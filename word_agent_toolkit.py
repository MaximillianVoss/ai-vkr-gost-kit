from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent
DEFAULT_SPEC_PATH = Path(__file__).resolve().with_name("ai-client-vkr-gost-spec.json")
WORD_FORMAT_DOCX = 16
WORD_FORMAT_PDF = 17
MSO_AUTOMATION_SECURITY_LOW = 1
MSO_AUTOMATION_SECURITY_FORCE_DISABLE = 3
WORD_PACKAGE_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml": "docx",
    "application/vnd.ms-word.document.macroEnabled.main+xml": "docm",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml": "dotx",
    "application/vnd.ms-word.template.macroEnabledTemplate.main+xml": "dotm",
}

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
WORD_ALIGNMENT_TO_COM = {
    "left": 0,
    "center": 1,
    "right": 2,
    "justify": 3,
}
WD_PARAGRAPH_UNIT = 4
WD_COLLAPSE_END = 0
WD_CAPTION_POSITION_BELOW = 0
WD_CAPTION_POSITION_ABOVE = 1
WD_ROW_ALIGN_CENTER = 1
WD_LINE_SPACE_SINGLE = 0
WD_LINE_SPACE_ONE_AND_HALF = 1
WD_LINE_SPACE_DOUBLE = 2
WORD_COLOR_BLACK = 0
BLACK_RGB = RGBColor(0, 0, 0)

HEADING_NUMBER_PREFIX_RE = re.compile(r"^\s*((\d+[\.\)\-]?\s*)+)([\t\s]*)", re.IGNORECASE)
APPENDIX_HEADING_RE = re.compile(r"^приложение(?:\s+([^\s]+))?(.*)$", re.IGNORECASE)
CAPTION_TEXT_PREFIXES = ("таблица ", "рисунок ", "рис. ")
SPECIAL_UNNUMBERED_TITLES = {
    "реферат": "Реферат",
    "введение": "Введение",
    "заключение": "Заключение",
    "список литературы": "Список литературы",
    "список использованных источников": "Список использованных источников",
    "список источников": "Список источников",
}

SAMPLE_VKR_CASES: list[dict[str, Any]] = [
    {
        "topic": "Разработка веб-платформы управления выпускными проектами",
        "domain": "процесс планирования, согласования и контроля этапов выпускных проектов на кафедре",
        "goal": "сокращение времени согласования материалов и повышение прозрачности контроля сроков",
        "object": "организация взаимодействия студентов, научных руководителей и методистов",
        "subject": "методы и программные средства автоматизации учебно-проектной деятельности",
        "pain_points": [
            "разрозненное хранение документов в почте и локальных каталогах",
            "отсутствие единого календаря контрольных точек и напоминаний",
            "сложность формирования сводной отчетности по группе и направлениям подготовки",
        ],
        "requirements": [
            "поддержка ролей студента, руководителя и методиста",
            "ведение графика этапов с фиксацией статусов и комментариев",
            "хранение версий документов и истории согласований",
            "экспорт отчетов по срокам, темам и проценту готовности",
        ],
        "entities": ["проект", "этап", "пользователь", "файл", "комментарий", "отчет"],
        "modules": [
            "модуль управления темами и карточками проектов",
            "модуль календарного планирования и уведомлений",
            "модуль согласования документов и хранения версий",
        ],
        "quality_metrics": [
            "время регистрации новой темы",
            "доля этапов, закрытых без просрочки",
            "скорость подготовки сводного отчета по группе",
        ],
        "effects": [
            "сокращение числа пропущенных контрольных точек",
            "уменьшение нагрузки на методиста при сборе статусов",
            "повышение качества коммуникации между участниками процесса",
        ],
        "sources": [
            "ГОСТ 7.32-2017",
            "учебный план направления подготовки",
            "локальный регламент подготовки выпускных работ",
            "материалы обследования кафедрального процесса",
            "документация к стеку Django и PostgreSQL",
        ],
        "appendix_items": [
            "карта пользовательских ролей и прав доступа",
            "пример карточки проекта",
            "фрагмент матрицы этапов и контрольных точек",
        ],
    },
    {
        "topic": "Проектирование сервиса аналитики обращений клиентов для малого бизнеса",
        "domain": "обработка входящих обращений из сайта, почты и мессенджеров в сервисной компании",
        "goal": "снижение времени первичного ответа и повышение качества управленческой аналитики",
        "object": "процесс регистрации, маршрутизации и анализа клиентских обращений",
        "subject": "модели агрегации, категоризации и визуализации операционных данных",
        "pain_points": [
            "повторный ручной ввод обращений из разных каналов",
            "неоднородные правила приоритизации и назначения исполнителей",
            "отсутствие наглядной аналитики по SLA, тематикам и загрузке сотрудников",
        ],
        "requirements": [
            "сбор обращений из нескольких каналов в едином интерфейсе",
            "автоматическое присвоение категории и приоритета обращения",
            "формирование панели показателей по срокам и повторным обращениям",
            "поддержка выгрузки аналитики для руководителя подразделения",
        ],
        "entities": ["обращение", "канал", "категория", "исполнитель", "SLA", "дашборд"],
        "modules": [
            "модуль интеграции с внешними каналами коммуникации",
            "модуль правил классификации и маршрутизации обращений",
            "модуль визуальной аналитики и мониторинга SLA",
        ],
        "quality_metrics": [
            "среднее время первого ответа",
            "доля обращений, обработанных в пределах SLA",
            "точность автоматической категоризации",
        ],
        "effects": [
            "ускорение реакции на обращения высокой критичности",
            "снижение числа потерянных заявок",
            "повышение обоснованности управленческих решений на основе данных",
        ],
        "sources": [
            "ГОСТ 7.32-2017",
            "регламент обработки клиентских обращений",
            "описание API подключаемых каналов",
            "методические материалы по проектированию BI-систем",
            "документация к FastAPI, Vue и ClickHouse",
        ],
        "appendix_items": [
            "схема потока обращений между каналами и сервисом",
            "пример дашборда руководителя",
            "таблица метрик качества обработки обращений",
        ],
    },
    {
        "topic": "Разработка информационной системы мониторинга энергопотребления учебного корпуса",
        "domain": "сбор и анализ показаний приборов учета для помещений учебного корпуса",
        "goal": "повышение точности контроля потребления ресурсов и поддержка мероприятий по энергосбережению",
        "object": "процесс наблюдения за динамикой энергопотребления помещений и инженерных зон",
        "subject": "инструменты централизованного мониторинга, хранения и анализа телеметрии",
        "pain_points": [
            "показания фиксируются в разное время и в разных форматах",
            "сложно выявлять отклонения и аномальные пики нагрузки",
            "нет единой визуальной картины по аудиториям, этажам и зонам корпуса",
        ],
        "requirements": [
            "регулярный прием показаний от приборов учета и операторов",
            "хранение истории измерений с детализацией по зонам и времени",
            "автоматическое выявление превышений пороговых значений",
            "предоставление отчетов по помещениям, этажам и категориям нагрузки",
        ],
        "entities": ["счетчик", "помещение", "зона", "измерение", "событие", "отчет"],
        "modules": [
            "модуль приема и валидации телеметрии",
            "модуль хранения временных рядов и событий отклонения",
            "модуль отчетности и визуального анализа динамики потребления",
        ],
        "quality_metrics": [
            "полнота поступления показаний по расписанию",
            "время обнаружения аномального потребления",
            "точность формирования месячного сводного отчета",
        ],
        "effects": [
            "ускорение реакции на нештатные режимы энергопотребления",
            "упрощение подготовки отчетов для административно-хозяйственной службы",
            "повышение эффективности планирования энергосберегающих мероприятий",
        ],
        "sources": [
            "ГОСТ 7.32-2017",
            "план помещений учебного корпуса",
            "регламент энергетического учета организации",
            "описание протоколов обмена с приборами учета",
            "документация к ASP.NET Core, InfluxDB и Grafana",
        ],
        "appendix_items": [
            "пример сводного отчета по этажам",
            "фрагмент журнала аномальных событий",
            "модель сущностей системы мониторинга",
        ],
    },
]


def load_spec(spec_path: Path | None) -> dict[str, Any]:
    current_path = spec_path or DEFAULT_SPEC_PATH
    with current_path.open("r", encoding="utf-8") as file:
        return json.load(file)


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def absolute_path(path: Path) -> str:
    return str(path.resolve())


def import_word_client():
    import win32com.client  # type: ignore[import]

    return win32com.client


def create_word_application(*, allow_macros: bool = False):
    client = import_word_client()
    word = client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    word.AutomationSecurity = MSO_AUTOMATION_SECURITY_LOW if allow_macros else MSO_AUTOMATION_SECURITY_FORCE_DISABLE
    if allow_macros:
        try:
            # Block AutoOpen/AutoExec and similar callbacks. Macros still run only when requested explicitly.
            word.WordBasic.DisableAutoMacros(1)
        except Exception:
            pass
    return word


def save_json(data: dict[str, Any], output_path: Path | None) -> None:
    payload = json.dumps(data, ensure_ascii=False, indent=2)
    if output_path is None:
        print(payload)
        return
    ensure_parent(output_path)
    output_path.write_text(payload, encoding="utf-8")


def normalize_lookup(value: str | None) -> str:
    if not value:
        return ""
    return " ".join(value.lower().replace("ё", "е").split())


def summarize_profile(profile: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": profile["id"],
        "title": profile["title"],
        "education_level": profile["education_level"],
        "qualification_level": profile["qualification_level"],
        "guideline_document": profile["guideline_document"],
        "when_to_use": profile["when_to_use"],
    }


def resolve_work_profile(
    spec: dict[str, Any],
    *,
    work_type: str | None,
    education_level: str | None,
    qualification_level: str | None,
) -> dict[str, Any]:
    routing = spec["work_type_routing"]
    normalized_input = normalize_lookup(" ".join(filter(None, [work_type, education_level, qualification_level])))
    matched_profiles: list[dict[str, Any]] = []

    for profile in routing["profiles"]:
        aliases = [normalize_lookup(alias) for alias in profile.get("aliases", [])]
        if any(alias and alias in normalized_input for alias in aliases):
            matched_profiles.append(profile)

    response: dict[str, Any] = {
        "input": {
            "work_type": work_type,
            "education_level": education_level,
            "qualification_level": qualification_level,
        },
        "selection_priority": routing["selection_priority"],
    }

    if len(matched_profiles) == 1:
        profile = matched_profiles[0]
        response.update(
            {
                "status": "resolved",
                "selected_profile": summarize_profile(profile),
                "selection_note": (
                    f"Выбран профиль {profile['id']} на основании совпадения типа работы, "
                    "уровня образования или квалификации."
                ),
            }
        )
        return response

    if len(matched_profiles) > 1:
        response.update(
            {
                "status": "ambiguous",
                "reason": "Найдено несколько подходящих профилей.",
                "candidates": [summarize_profile(profile) for profile in matched_profiles],
                "recommended_action": routing["ambiguity_rules"][1]["action"],
            }
        )
        return response

    if "вкр" in normalized_input:
        response.update(
            {
                "status": "ambiguous",
                "reason": "Слово «ВКР» само по себе неоднозначно и не определяет уровень образования.",
                "candidates": [summarize_profile(profile) for profile in routing["profiles"]],
                "recommended_action": routing["ambiguity_rules"][0]["action"],
            }
        )
        return response

    response.update(
        {
            "status": "unresolved",
            "reason": "Не удалось однозначно сопоставить входные данные с профилями методических указаний.",
            "recommended_action": "Уточнить тип работы, уровень образования и основную методичку.",
            "available_profiles": [summarize_profile(profile) for profile in routing["profiles"]],
        }
    )
    return response


def inspect_word_package(document_path: Path) -> dict[str, Any]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(document_path) as archive:
        content_types_xml = archive.read("[Content_Types].xml")
        styles_xml = archive.read("word/styles.xml")
        content_types_root = ET.fromstring(content_types_xml)
        vba_parts = sorted(entry.filename for entry in archive.infolist() if entry.filename.endswith("vbaProject.bin"))
    root = ET.fromstring(styles_xml)
    package_content_type = ""

    for override in content_types_root.findall("{http://schemas.openxmlformats.org/package/2006/content-types}Override"):
        if override.get("PartName") == "/word/document.xml":
            package_content_type = override.get("ContentType", "")
            break

    styles_by_type: dict[str, list[dict[str, str]]] = {
        "paragraph": [],
        "character": [],
        "table": [],
        "numbering": [],
    }

    for style in root.findall("w:style", ns):
        style_type = style.get(f"{{{ns['w']}}}type", "unknown")
        style_id = style.get(f"{{{ns['w']}}}styleId", "")
        name_element = style.find("w:name", ns)
        style_name = name_element.get(f"{{{ns['w']}}}val", style_id) if name_element is not None else style_id
        target = styles_by_type.setdefault(style_type, [])
        target.append({"style_id": style_id, "name": style_name})

    for values in styles_by_type.values():
        values.sort(key=lambda item: item["name"].lower())

    return {
        "document_path": str(document_path),
        "package_type": WORD_PACKAGE_TYPES.get(package_content_type, "unknown"),
        "main_content_type": package_content_type,
        "has_vba_project": bool(vba_parts),
        "vba_parts": vba_parts,
        "styles": styles_by_type,
    }


def iter_macro_candidates(document, macro_name: str) -> list[str]:
    if "!" in macro_name:
        return [macro_name]

    containers: list[str] = []
    for value in [getattr(document, "FullName", None), getattr(document, "Name", None)]:
        if value:
            containers.append(str(value))

    try:
        attached_template = document.AttachedTemplate
    except Exception:
        attached_template = None

    for value in [
        getattr(attached_template, "FullName", None),
        getattr(attached_template, "Name", None),
        attached_template,
    ]:
        if value:
            containers.append(str(value))

    candidates = [macro_name]
    seen: set[str] = set()
    for container in containers:
        parts = [container]
        try:
            container_path = Path(container)
        except (TypeError, ValueError):
            container_path = None
        if container_path is not None:
            parts.extend([container_path.name, container_path.stem])
        for part in parts:
            if not part:
                continue
            qualified = f"'{part}'!{macro_name}"
            if qualified in seen:
                continue
            seen.add(qualified)
            candidates.append(qualified)

    return candidates


def run_requested_macros(word, document, macro_names: list[str] | None) -> None:
    if not macro_names:
        return

    for macro_name in macro_names:
        attempts: list[str] = []
        last_error: Exception | None = None
        for candidate in iter_macro_candidates(document, macro_name):
            attempts.append(candidate)
            try:
                word.Run(candidate)
                last_error = None
                break
            except Exception as error:
                last_error = error
        if last_error is not None:
            attempts_text = ", ".join(attempts)
            raise RuntimeError(
                f"Не удалось выполнить макрос '{macro_name}'. Проверены варианты: {attempts_text}."
            ) from last_error


def update_document_fields(document) -> None:
    document.Fields.Update()
    for toc in document.TablesOfContents:
        toc.Update()


def unlink_toc_fields(document) -> None:
    for toc in document.TablesOfContents:
        try:
            fields = toc.Range.Fields
        except Exception:
            continue
        try:
            count = int(fields.Count)
        except Exception:
            count = 0
        for index in range(count, 0, -1):
            try:
                fields(index).Unlink()
            except Exception:
                continue


def clean_word_text(value: str) -> str:
    return value.replace("\r", "").replace("\x07", "").strip()


def apply_text_case(value: str, text_case: str | None) -> str:
    if text_case == "uppercase":
        return value.upper()
    if text_case == "lowercase":
        return value.lower()
    return value


def get_word_style_name(value: Any) -> str:
    for attribute in ["NameLocal", "Name"]:
        try:
            candidate = getattr(value, attribute)
        except Exception:
            continue
        if candidate:
            return str(candidate)
    return str(value)


def get_word_paragraph_style_name(paragraph) -> str:
    return get_word_style_name(paragraph.Style)


def find_word_style_name(document, candidates: list[str]) -> str | None:
    for candidate in candidates:
        try:
            style = document.Styles(candidate)
            return get_word_style_name(style)
        except Exception:
            continue
    return None


def set_word_line_spacing(paragraph_format, line_spacing: float) -> None:
    if line_spacing == 1.0:
        paragraph_format.LineSpacingRule = WD_LINE_SPACE_SINGLE
        return
    if line_spacing == 1.5:
        paragraph_format.LineSpacingRule = WD_LINE_SPACE_ONE_AND_HALF
        return
    if line_spacing == 2.0:
        paragraph_format.LineSpacingRule = WD_LINE_SPACE_DOUBLE
        return


def format_word_paragraph(
    paragraph,
    *,
    font_name: str,
    font_size_pt: float,
    bold: bool | None = None,
    alignment: str | None = None,
    line_spacing: float | None = None,
    first_line_indent_points: float | None = None,
    space_before_points: float | None = None,
    space_after_points: float | None = None,
) -> None:
    paragraph.Range.Font.Name = font_name
    paragraph.Range.Font.Size = font_size_pt
    paragraph.Range.Font.Color = WORD_COLOR_BLACK
    if bold is not None:
        paragraph.Range.Font.Bold = -1 if bold else 0
    paragraph_format = paragraph.Range.ParagraphFormat
    if alignment is not None:
        paragraph_format.Alignment = WORD_ALIGNMENT_TO_COM[alignment]
    if line_spacing is not None:
        set_word_line_spacing(paragraph_format, line_spacing)
    if first_line_indent_points is not None:
        paragraph_format.FirstLineIndent = first_line_indent_points
    paragraph_format.SpaceBefore = 0 if space_before_points is None else space_before_points
    paragraph_format.SpaceAfter = 0 if space_after_points is None else space_after_points


def format_caption_paragraph(paragraph, *, font_name: str, font_size_pt: float, alignment: str) -> None:
    format_word_paragraph(
        paragraph,
        font_name=font_name,
        font_size_pt=font_size_pt,
        alignment=alignment,
        line_spacing=1.0,
        first_line_indent_points=0,
    )
    paragraph.Range.ParagraphFormat.LeftIndent = 0
    paragraph.Range.ParagraphFormat.RightIndent = 0


def looks_like_caption(text: str) -> bool:
    normalized_text = normalize_lookup(text)
    return normalized_text.startswith(CAPTION_TEXT_PREFIXES)


def canonicalize_appendix_heading(text: str) -> str:
    match = APPENDIX_HEADING_RE.match(text.strip())
    if match is None:
        return text
    label = (match.group(1) or "").strip()
    suffix = match.group(2) or ""
    parts = ["Приложение"]
    if label:
        parts.append(label.upper())
    return " ".join(parts) + suffix


def get_chapter_text_case(spec: dict[str, Any]) -> str | None:
    return spec["heading_rules"]["chapter_title"].get("text_case")


def remove_word_paragraph_numbering(paragraph) -> None:
    try:
        paragraph.Range.ListFormat.RemoveNumbers()
    except Exception:
        pass


def normalize_special_heading_titles(document, spec: dict[str, Any], *, strip_heading_numbering: bool = False) -> None:
    style_map = spec["template_style_map"]
    heading_style_names = {
        normalize_lookup(name)
        for name in (
            list(style_map["chapter_title"]) + list(style_map["section_title"]) + list(style_map["subsection_title"])
        )
    }
    chapter_text_case = get_chapter_text_case(spec)

    for paragraph in document.Paragraphs:
        original_style_name = get_word_style_name(paragraph.Style)
        style_name = normalize_lookup(original_style_name)
        if style_name not in heading_style_names:
            continue

        text = clean_word_text(paragraph.Range.Text)
        if not text:
            continue

        stripped_text = HEADING_NUMBER_PREFIX_RE.sub("", text).strip()
        normalized_stripped = normalize_lookup(stripped_text)
        content_range = paragraph.Range.Duplicate
        if content_range.End > content_range.Start:
            content_range.End -= 1

        if normalized_stripped in SPECIAL_UNNUMBERED_TITLES:
            content_range.Text = apply_text_case(SPECIAL_UNNUMBERED_TITLES[normalized_stripped], chapter_text_case)
            remove_word_paragraph_numbering(paragraph)
            try:
                paragraph.Style = original_style_name
            except Exception:
                pass
            paragraph.Range.ParagraphFormat.Alignment = WORD_ALIGNMENT_TO_COM["center"]
            continue

        if normalized_stripped.startswith("приложение"):
            content_range.Text = apply_text_case(canonicalize_appendix_heading(stripped_text), chapter_text_case)
            remove_word_paragraph_numbering(paragraph)
            try:
                paragraph.Style = original_style_name
            except Exception:
                pass
            paragraph.Range.ParagraphFormat.Alignment = WORD_ALIGNMENT_TO_COM["center"]
            continue

        if strip_heading_numbering:
            remove_word_paragraph_numbering(paragraph)


def format_heading_paragraphs(document, spec: dict[str, Any]) -> None:
    style_map = spec["template_style_map"]
    heading_rules = [
        (list(style_map["chapter_title"]), spec["heading_rules"]["chapter_title"], 1),
        (list(style_map["section_title"]), spec["heading_rules"]["section_title"], 2),
        (list(style_map["subsection_title"]), spec["heading_rules"]["subsection_title"], 3),
    ]
    normalized_rules = {
        normalize_lookup(style_name): (rules, outline_level)
        for style_names, rules, outline_level in heading_rules
        for style_name in style_names
    }

    for paragraph in document.Paragraphs:
        text = clean_word_text(paragraph.Range.Text)
        if not text:
            continue

        style_name = normalize_lookup(get_word_paragraph_style_name(paragraph))
        entry = normalized_rules.get(style_name)
        if entry is None:
            continue
        rules, outline_level = entry

        format_word_paragraph(
            paragraph,
            font_name=rules["font_name"],
            font_size_pt=rules["font_size_pt"],
            bold=rules["bold"],
            alignment=rules["alignment"],
            line_spacing=rules["multi_line_spacing"],
            first_line_indent_points=0,
            space_before_points=rules.get("space_before_pt"),
            space_after_points=rules.get("space_after_pt"),
        )
        try:
            paragraph.OutlineLevel = outline_level
        except Exception:
            pass


def format_toc_paragraphs(document, spec: dict[str, Any]) -> None:
    style_map = spec["template_style_map"]
    toc_title_style_names = {normalize_lookup(name) for name in style_map["toc_heading"]}
    toc_entry_style_names = {normalize_lookup(name) for name in style_map["toc_entry"]}
    toc_entry_style_names.update({normalize_lookup(f"TOC {index}") for index in range(1, 10)})
    toc_entry_style_names.update({normalize_lookup(f"toc {index}") for index in range(1, 10)})

    toc_title_rules = spec["toc_rules"]["title_format"]
    toc_entry_rules = spec["toc_rules"]["entry_format"]

    for paragraph in document.Paragraphs:
        text = clean_word_text(paragraph.Range.Text)
        if not text:
            continue

        style_name = normalize_lookup(get_word_paragraph_style_name(paragraph))
        if style_name in toc_title_style_names:
            remove_word_paragraph_numbering(paragraph)
            format_word_paragraph(
                paragraph,
                font_name=toc_title_rules["font_name"],
                font_size_pt=toc_title_rules["font_size_pt"],
                bold=toc_title_rules["bold"],
                alignment=toc_title_rules["alignment"],
                line_spacing=toc_title_rules["line_spacing"],
                first_line_indent_points=0,
            )
            continue

        if style_name in toc_entry_style_names:
            parts = text.split("\t")
            if len(parts) >= 3 and re.fullmatch(r"\d+(?:\.\d+)*", parts[0].strip()):
                title = "\t".join(parts[1:-1]).strip()
                page = parts[-1].strip()
                normalized_text = f"{title}\t{page}"
                content_range = paragraph.Range.Duplicate
                if content_range.End > content_range.Start:
                    content_range.End -= 1
                content_range.Text = normalized_text
            format_word_paragraph(
                paragraph,
                font_name=toc_entry_rules["font_name"],
                font_size_pt=toc_entry_rules["font_size_pt"],
                alignment=toc_entry_rules["alignment"],
                line_spacing=toc_entry_rules["line_spacing"],
                first_line_indent_points=0,
            )


def format_body_paragraphs_without_headings(document, spec: dict[str, Any]) -> None:
    style_map = spec["template_style_map"]
    body_rules = spec["global_formatting_rules"]["main_text"]
    excluded_style_names = {
        normalize_lookup(name)
        for name in (
            list(style_map["chapter_title"])
            + list(style_map["section_title"])
            + list(style_map["subsection_title"])
            + list(style_map["toc_heading"])
            + list(style_map["toc_entry"])
            + list(style_map["caption"])
            + list(style_map["header"])
            + list(style_map["footer"])
        )
    }

    for paragraph in document.Paragraphs:
        text = clean_word_text(paragraph.Range.Text)
        if not text or looks_like_caption(text):
            continue

        style_name = normalize_lookup(get_word_paragraph_style_name(paragraph))
        if style_name in excluded_style_names:
            continue

        paragraph.Range.Font.Name = body_rules["font_name"]
        paragraph.Range.Font.Size = body_rules["font_size_pt"]
        paragraph.Range.Font.Color = WORD_COLOR_BLACK
        set_word_line_spacing(paragraph.Range.ParagraphFormat, body_rules["line_spacing"])
        paragraph.Range.ParagraphFormat.SpaceBefore = 0
        paragraph.Range.ParagraphFormat.SpaceAfter = 0


def get_previous_paragraph(item) -> Any | None:
    try:
        previous_range = item.Range.Previous(WD_PARAGRAPH_UNIT)
    except Exception:
        return None
    if previous_range is None:
        return None
    try:
        return previous_range.Paragraphs(1)
    except Exception:
        return None


def get_next_paragraph(item) -> Any | None:
    try:
        next_range = item.Range.Next(WD_PARAGRAPH_UNIT)
    except Exception:
        return None
    if next_range is None:
        return None
    try:
        return next_range.Paragraphs(1)
    except Exception:
        return None


def has_table_caption(table, caption_style_name: str | None) -> bool:
    paragraph = get_previous_paragraph(table)
    if paragraph is None:
        return False
    text = clean_word_text(paragraph.Range.Text)
    if not text:
        return False
    if looks_like_caption(text):
        return True
    if caption_style_name is None:
        return False
    return normalize_lookup(get_word_paragraph_style_name(paragraph)) == normalize_lookup(caption_style_name)


def has_figure_caption(shape, caption_style_name: str | None) -> bool:
    paragraph = get_next_paragraph(shape)
    if paragraph is None:
        return False
    text = clean_word_text(paragraph.Range.Text)
    if not text:
        return False
    if looks_like_caption(text):
        return True
    if caption_style_name is None:
        return False
    return normalize_lookup(get_word_paragraph_style_name(paragraph)) == normalize_lookup(caption_style_name)


def format_tables(document, spec: dict[str, Any], *, insert_placeholder_captions: bool = False) -> None:
    table_rules = spec["table_rules"]
    table_font = table_rules["table_font"]
    caption_style_name = find_word_style_name(document, list(spec["template_style_map"]["caption"]))
    table_style_name = find_word_style_name(document, ["ГОСТ.Черный", "ГОСТ.Синий", "Table Grid"])
    selection = document.Application.Selection

    for table in document.Tables:
        if table_style_name is not None:
            try:
                table.Style = table_style_name
            except Exception:
                pass

        try:
            table.Rows.Alignment = WD_ROW_ALIGN_CENTER
        except Exception:
            pass

        for row_index in range(1, table.Rows.Count + 1):
            row = table.Rows(row_index)
            cell_alignment = "center" if row_index == 1 else table_rules["alignment"]["body_cells"]
            for cell in row.Cells:
                for paragraph in cell.Range.Paragraphs:
                    text = clean_word_text(paragraph.Range.Text)
                    if not text:
                        continue
                    format_word_paragraph(
                        paragraph,
                        font_name=table_font["font_name"],
                        font_size_pt=table_font["font_size_pt"],
                        alignment=cell_alignment,
                        line_spacing=1.0,
                        first_line_indent_points=0,
                    )
                    paragraph.Range.ParagraphFormat.LeftIndent = 0
                    paragraph.Range.ParagraphFormat.RightIndent = 0

        if not insert_placeholder_captions or has_table_caption(table, caption_style_name):
            continue

        table.Range.Select()
        selection.InsertCaption(Label="Таблица", Title=" Описание", Position=WD_CAPTION_POSITION_ABOVE)
        inserted_paragraph = selection.Paragraphs(1)
        if caption_style_name is not None:
            try:
                inserted_paragraph.Style = caption_style_name
            except Exception:
                pass
        format_caption_paragraph(
            inserted_paragraph,
            font_name=table_font["font_name"],
            font_size_pt=table_font["font_size_pt"],
            alignment=table_rules["caption_layout"]["number_alignment"],
        )


def format_inline_shapes(document, spec: dict[str, Any], *, insert_placeholder_captions: bool = False) -> None:
    figure_rules = spec["figure_rules"]
    table_font = spec["table_rules"]["table_font"]
    caption_style_name = find_word_style_name(document, list(spec["template_style_map"]["caption"]))
    selection = document.Application.Selection

    for inline_shape in document.InlineShapes:
        paragraph = inline_shape.Range.Paragraphs(1)
        paragraph.Range.ParagraphFormat.Alignment = WORD_ALIGNMENT_TO_COM[figure_rules["alignment"]]
        paragraph.Range.ParagraphFormat.LeftIndent = 0
        paragraph.Range.ParagraphFormat.RightIndent = 0
        paragraph.Range.ParagraphFormat.FirstLineIndent = 0

        try:
            section = inline_shape.Range.Sections(1)
            page_width = section.PageSetup.PageWidth - section.PageSetup.LeftMargin - section.PageSetup.RightMargin
            if inline_shape.Width > page_width:
                inline_shape.LockAspectRatio = True
                inline_shape.Width = page_width
        except Exception:
            pass

        if not insert_placeholder_captions or has_figure_caption(inline_shape, caption_style_name):
            continue

        inline_shape.Range.Select()
        selection.Collapse(WD_COLLAPSE_END)
        selection.InsertParagraphAfter()
        selection.Collapse(WD_COLLAPSE_END)
        selection.InsertCaption(Label="Рисунок", Title=" Описание", Position=WD_CAPTION_POSITION_BELOW)
        inserted_paragraph = selection.Paragraphs(1)
        if caption_style_name is not None:
            try:
                inserted_paragraph.Style = caption_style_name
            except Exception:
                pass
        format_caption_paragraph(
            inserted_paragraph,
            font_name=table_font["font_name"],
            font_size_pt=table_font["font_size_pt"],
            alignment=figure_rules["caption_alignment"],
        )


def apply_builtin_layout_rules(
    document,
    spec: dict[str, Any],
    *,
    insert_placeholder_captions: bool = False,
    strip_heading_numbering: bool = False,
) -> None:
    normalize_special_heading_titles(document, spec, strip_heading_numbering=strip_heading_numbering)
    format_heading_paragraphs(document, spec)
    format_tables(document, spec, insert_placeholder_captions=insert_placeholder_captions)
    format_inline_shapes(document, spec, insert_placeholder_captions=insert_placeholder_captions)
    format_body_paragraphs_without_headings(document, spec)
    format_toc_paragraphs(document, spec)


def process_word_document(
    docx_path: Path,
    *,
    pdf_path: Path | None = None,
    macro_names: list[str] | None = None,
    spec_path: Path | None = None,
    builtin_layout: bool = False,
    insert_placeholder_captions: bool = False,
    strip_heading_numbering: bool = False,
) -> None:
    if pdf_path is not None:
        ensure_parent(pdf_path)

    spec = load_spec(spec_path) if builtin_layout else None
    word = create_word_application(allow_macros=bool(macro_names))
    document = None
    try:
        document = word.Documents.Open(absolute_path(docx_path))
        run_requested_macros(word, document, macro_names)
        if spec is not None:
            apply_builtin_layout_rules(
                document,
                spec,
                insert_placeholder_captions=insert_placeholder_captions,
                strip_heading_numbering=strip_heading_numbering,
            )
        update_document_fields(document)
        if spec is not None:
            apply_builtin_layout_rules(
                document,
                spec,
                insert_placeholder_captions=insert_placeholder_captions,
                strip_heading_numbering=strip_heading_numbering,
            )
        document.Save()
        if pdf_path is not None:
            if spec is not None:
                unlink_toc_fields(document)
            document.ExportAsFixedFormat(absolute_path(pdf_path), WORD_FORMAT_PDF)
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()


def create_docx_from_template(template_path: Path, output_path: Path, *, macro_names: list[str] | None = None) -> None:
    ensure_parent(output_path)
    word = create_word_application(allow_macros=bool(macro_names))
    document = None
    try:
        document = word.Documents.Add(absolute_path(template_path))
        run_requested_macros(word, document, macro_names)
        document.SaveAs(absolute_path(output_path), FileFormat=WORD_FORMAT_DOCX)
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()


def refresh_fields(
    docx_path: Path,
    *,
    macro_names: list[str] | None = None,
    spec_path: Path | None = None,
    builtin_layout: bool = False,
    insert_placeholder_captions: bool = False,
    strip_heading_numbering: bool = False,
) -> None:
    process_word_document(
        docx_path,
        macro_names=macro_names,
        spec_path=spec_path,
        builtin_layout=builtin_layout,
        insert_placeholder_captions=insert_placeholder_captions,
        strip_heading_numbering=strip_heading_numbering,
    )


def export_pdf(
    docx_path: Path,
    pdf_path: Path,
    *,
    macro_names: list[str] | None = None,
    spec_path: Path | None = None,
    builtin_layout: bool = False,
    insert_placeholder_captions: bool = False,
    strip_heading_numbering: bool = False,
) -> None:
    process_word_document(
        docx_path,
        pdf_path=pdf_path,
        macro_names=macro_names,
        spec_path=spec_path,
        builtin_layout=builtin_layout,
        insert_placeholder_captions=insert_placeholder_captions,
        strip_heading_numbering=strip_heading_numbering,
    )


def get_style_name(document: Document, candidates: list[str], fallback: str) -> str:
    existing_names = {style.name for style in document.styles}
    for candidate in candidates:
        if candidate in existing_names:
            return candidate
    return fallback


def get_existing_style(document: Document, candidates: list[str]) -> Any | None:
    existing_names = {style.name for style in document.styles}
    for candidate in candidates:
        if candidate in existing_names:
            return document.styles[candidate]
    return None


def get_existing_styles(document: Document, candidates: list[str]) -> list[Any]:
    existing_names = {style.name for style in document.styles}
    return [document.styles[candidate] for candidate in candidates if candidate in existing_names]


def clear_document(document: Document) -> None:
    body = document.element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def ensure_font_family(style, font_name: str) -> None:
    style.font.name = font_name
    rpr = style._element.get_or_add_rPr()
    if rpr.rFonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rpr.rFonts.set(qn("w:ascii"), font_name)
    rpr.rFonts.set(qn("w:hAnsi"), font_name)
    rpr.rFonts.set(qn("w:eastAsia"), font_name)
    rpr.rFonts.set(qn("w:cs"), font_name)


def rebase_style(style, base_style) -> None:
    try:
        style.base_style = base_style
    except Exception:
        pass

    based_on = None
    for child in style._element:
        if child.tag == qn("w:basedOn"):
            based_on = child
            break
    if base_style is None:
        if based_on is not None:
            style._element.remove(based_on)
        return
    if based_on is None:
        based_on = OxmlElement("w:basedOn")
        style._element.insert(1, based_on)
    based_on.set(qn("w:val"), base_style.style_id)


def clear_style_numbering(style) -> None:
    paragraph_properties = style._element.find(qn("w:pPr"))
    if paragraph_properties is None:
        return
    for child in list(paragraph_properties):
        if child.tag == qn("w:numPr"):
            paragraph_properties.remove(child)


def apply_paragraph_style(
    style,
    *,
    font_name: str,
    font_size_pt: float,
    bold: bool | None = None,
    alignment: str | None = None,
    line_spacing: float | None = None,
    first_line_indent_cm: float | None = None,
    space_before_pt: float | None = None,
    space_after_pt: float | None = None,
) -> None:
    clear_style_numbering(style)
    ensure_font_family(style, font_name)
    style.font.size = Pt(font_size_pt)
    style.font.color.rgb = BLACK_RGB
    if bold is not None:
        style.font.bold = bold

    paragraph_format = style.paragraph_format
    if alignment is not None:
        paragraph_format.alignment = ALIGNMENTS[alignment]
    if line_spacing is not None:
        paragraph_format.line_spacing = line_spacing
    if first_line_indent_cm is None:
        paragraph_format.first_line_indent = None
    else:
        paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    paragraph_format.space_before = Pt(0 if space_before_pt is None else space_before_pt)
    paragraph_format.space_after = Pt(0 if space_after_pt is None else space_after_pt)


def apply_run_font(
    run,
    *,
    font_name: str = "Times New Roman",
    font_size_pt: float | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = font_name
    run.font.color.rgb = BLACK_RGB
    rpr = run._element.get_or_add_rPr()
    if rpr.rFonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rpr.rFonts.set(qn("w:ascii"), font_name)
    rpr.rFonts.set(qn("w:hAnsi"), font_name)
    rpr.rFonts.set(qn("w:eastAsia"), font_name)
    rpr.rFonts.set(qn("w:cs"), font_name)
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
    if bold is not None:
        run.bold = bold


def add_field(
    paragraph,
    instruction: str,
    placeholder: str = "",
    *,
    font_name: str = "Times New Roman",
    font_size_pt: float = 12,
    bold: bool | None = None,
) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    paragraph.add_run()._r.append(begin)
    paragraph.add_run()._r.append(instr)
    paragraph.add_run()._r.append(separate)
    if placeholder:
        apply_run_font(
            paragraph.add_run(placeholder),
            font_name=font_name,
            font_size_pt=font_size_pt,
            bold=bold,
        )
    paragraph.add_run()._r.append(end)


def set_outline_level(paragraph, level: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    for child in list(paragraph_properties):
        if child.tag == qn("w:outlineLvl"):
            paragraph_properties.remove(child)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    paragraph_properties.append(outline)


def add_paragraph(
    document: Document,
    style_name: str,
    text: str = "",
    *,
    alignment: str = "justify",
    first_line_indent_cm: float | None = 1.25,
    bold: bool = False,
    font_size_pt: float | None = None,
    outline_level: int | None = None,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.style = style_name
    paragraph.alignment = ALIGNMENTS[alignment]
    if first_line_indent_cm is None:
        paragraph.paragraph_format.first_line_indent = None
    else:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    if outline_level is not None:
        set_outline_level(paragraph, outline_level)
    if text:
        apply_run_font(paragraph.add_run(text), font_size_pt=font_size_pt, bold=bold)
    return paragraph


def add_blank(document: Document, style_name: str, count: int = 1) -> None:
    for _ in range(count):
        add_paragraph(document, style_name, alignment="left", first_line_indent_cm=None)


def add_list_item(
    document: Document,
    style_name: str,
    text: str,
    *,
    marker: str = "-",
    level: int = 0,
) -> Any:
    paragraph = add_paragraph(
        document,
        style_name,
        f"{marker} {text}",
        alignment="left",
        first_line_indent_cm=None,
    )
    paragraph.paragraph_format.left_indent = Cm(0.75 * level)
    return paragraph


def add_heading(
    document: Document,
    style_name: str,
    text: str,
    *,
    outline_level: int,
    alignment: str,
    page_break: bool = False,
) -> Any:
    if page_break:
        document.add_page_break()
    return add_paragraph(
        document,
        style_name,
        text,
        alignment=alignment,
        first_line_indent_cm=None,
        bold=True,
        outline_level=outline_level,
    )


def clear_story_container(container) -> None:
    for child in list(container._element):
        container._element.remove(child)


def create_footer_paragraph(section, style_name: str, spec: dict[str, Any]):
    page_rules = spec["global_formatting_rules"]["page_numbering"]
    paragraph = section.footer.add_paragraph()
    paragraph.style = style_name
    paragraph.alignment = ALIGNMENTS["right" if page_rules["position"] == "bottom-right" else "center"]
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    add_field(
        paragraph,
        "PAGE",
        "1",
        font_name=page_rules["font_name"],
        font_size_pt=page_rules["font_size_pt"],
    )
    return paragraph


def configure_page_numbering_sections(document: Document, spec: dict[str, Any]) -> None:
    page_rules = spec["global_formatting_rules"]["page_numbering"]
    style_map = spec["template_style_map"]
    footer_style = get_style_name(document, style_map["footer"], "Footer")
    sections = list(document.sections)
    toc_section_index = 1 if len(sections) >= 3 else None

    for index, section in enumerate(sections):
        section.footer.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        section.different_first_page_header_footer = False

        clear_story_container(section.footer)
        clear_story_container(section.first_page_footer)

        is_title_section = index == 0
        is_toc_section = toc_section_index is not None and index == toc_section_index
        show_page_number = page_rules["print_on_other_pages"]
        if is_title_section:
            show_page_number = page_rules["print_on_title_page"]
        elif is_toc_section:
            show_page_number = page_rules["print_on_toc_page"]

        if show_page_number:
            create_footer_paragraph(section, footer_style, spec)


def build_toc_instruction(spec: dict[str, Any]) -> str:
    return r'TOC \o "1-3" \z \u'


def sanitize_filename_component(value: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]+', "", value)
    sanitized = re.sub(r"\s+", " ", sanitized).strip()
    sanitized = sanitized.replace(" ", "_")
    return sanitized[:80] or "document"


def set_default_section(section, spec: dict[str, Any]) -> None:
    margins = spec["global_formatting_rules"]["margins_mm"]
    section.left_margin = Cm(margins["left"] / 10)
    section.right_margin = Cm(margins["right"] / 10)
    section.top_margin = Cm(margins["top"] / 10)
    section.bottom_margin = Cm(margins["bottom"] / 10)
    section.header_distance = Cm(spec["global_formatting_rules"]["header_distance_cm"])
    section.footer_distance = Cm(spec["global_formatting_rules"]["footer_distance_cm"])


def apply_gost_profile(input_path: Path, output_path: Path | None, spec_path: Path | None) -> Path:
    spec = load_spec(spec_path)
    target_path = output_path or input_path
    document = Document(str(input_path))

    style_map = spec["template_style_map"]
    body_rules = spec["global_formatting_rules"]["main_text"]
    body_style_name = get_style_name(document, style_map["body"], "Normal")
    body_style = document.styles[body_style_name]
    apply_paragraph_style(
        body_style,
        font_name=body_rules["font_name"],
        font_size_pt=body_rules["font_size_pt"],
        alignment=body_rules["alignment"],
        line_spacing=body_rules["line_spacing"],
        first_line_indent_cm=body_rules["first_line_indent_cm"],
    )

    chapter_rules = spec["heading_rules"]["chapter_title"]
    for chapter_style in get_existing_styles(document, style_map["chapter_title"]):
        if chapter_style.name.startswith("ГОСТ "):
            rebase_style(chapter_style, body_style)
        apply_paragraph_style(
            chapter_style,
            font_name=chapter_rules["font_name"],
            font_size_pt=chapter_rules["font_size_pt"],
            bold=chapter_rules["bold"],
            alignment=chapter_rules["alignment"],
            line_spacing=chapter_rules["multi_line_spacing"],
            first_line_indent_cm=None,
            space_before_pt=chapter_rules.get("space_before_pt"),
            space_after_pt=chapter_rules.get("space_after_pt"),
        )

    section_rules = spec["heading_rules"]["section_title"]
    for section_style in get_existing_styles(document, style_map["section_title"]):
        if section_style.name.startswith("ГОСТ "):
            rebase_style(section_style, body_style)
        apply_paragraph_style(
            section_style,
            font_name=section_rules["font_name"],
            font_size_pt=section_rules["font_size_pt"],
            bold=section_rules["bold"],
            alignment=section_rules["alignment"],
            line_spacing=section_rules["multi_line_spacing"],
            first_line_indent_cm=None,
            space_before_pt=section_rules.get("space_before_pt"),
            space_after_pt=section_rules.get("space_after_pt"),
        )

    subsection_rules = spec["heading_rules"]["subsection_title"]
    for subsection_style in get_existing_styles(document, style_map["subsection_title"]):
        if subsection_style.name.startswith("ГОСТ "):
            rebase_style(subsection_style, body_style)
        apply_paragraph_style(
            subsection_style,
            font_name=subsection_rules["font_name"],
            font_size_pt=subsection_rules["font_size_pt"],
            bold=subsection_rules["bold"],
            alignment=subsection_rules["alignment"],
            line_spacing=subsection_rules["multi_line_spacing"],
            first_line_indent_cm=None,
            space_before_pt=subsection_rules.get("space_before_pt"),
            space_after_pt=subsection_rules.get("space_after_pt"),
        )

    for caption_style in get_existing_styles(document, style_map["caption"]):
        apply_paragraph_style(
            caption_style,
            font_name=spec["table_rules"]["table_font"]["font_name"],
            font_size_pt=12,
            alignment="center",
            line_spacing=1.0,
            first_line_indent_cm=None,
        )

    for toc_title_style in get_existing_styles(document, style_map["toc_heading"]):
        rebase_style(toc_title_style, body_style)
        toc_title_rules = spec["toc_rules"]["title_format"]
        apply_paragraph_style(
            toc_title_style,
            font_name=toc_title_rules["font_name"],
            font_size_pt=toc_title_rules["font_size_pt"],
            bold=toc_title_rules["bold"],
            alignment=toc_title_rules["alignment"],
            line_spacing=toc_title_rules["line_spacing"],
            first_line_indent_cm=None,
        )

    toc_entry_candidates = list(style_map["toc_entry"])
    toc_entry_candidates.extend([f"TOC {index}" for index in range(1, 10)])
    toc_entry_candidates.extend([f"toc {index}" for index in range(1, 10)])
    toc_entry_rules = spec["toc_rules"]["entry_format"]
    applied_toc_styles: set[str] = set()
    for candidate in toc_entry_candidates:
        if candidate in applied_toc_styles:
            continue
        style = get_existing_style(document, [candidate])
        if style is None:
            continue
        apply_paragraph_style(
            style,
            font_name=toc_entry_rules["font_name"],
            font_size_pt=toc_entry_rules["font_size_pt"],
            alignment=toc_entry_rules["alignment"],
            line_spacing=toc_entry_rules["line_spacing"],
            first_line_indent_cm=toc_entry_rules["first_line_indent_cm"],
        )
        applied_toc_styles.add(candidate)

    for footer_style in get_existing_styles(document, style_map["footer"]):
        page_rules = spec["global_formatting_rules"]["page_numbering"]
        apply_paragraph_style(
            footer_style,
            font_name=page_rules["font_name"],
            font_size_pt=page_rules["font_size_pt"],
            alignment="right" if page_rules["position"] == "bottom-right" else "center",
            line_spacing=1.0,
            first_line_indent_cm=None,
        )

    for section in document.sections:
        set_default_section(section, spec)

    ensure_parent(target_path)
    document.save(target_path)
    return target_path


def append_code_appendix(
    input_path: Path,
    source_path: Path,
    title: str,
    appendix_label: str,
    output_path: Path | None,
    spec_path: Path | None,
    line_numbers: bool,
) -> Path:
    spec = load_spec(spec_path)
    target_path = output_path or input_path
    document = Document(str(input_path))
    style_map = spec["template_style_map"]
    body_style_name = get_style_name(document, style_map["body"], "Normal")

    document.add_page_break()

    appendix_paragraph = document.add_paragraph()
    appendix_paragraph.style = body_style_name
    appendix_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    appendix_paragraph.paragraph_format.first_line_indent = None
    appendix_paragraph.add_run(f"Приложение {appendix_label}").bold = True

    title_paragraph = document.add_paragraph()
    title_paragraph.style = body_style_name
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.first_line_indent = None
    title_paragraph.add_run(title).bold = True

    document.add_paragraph().style = body_style_name

    code_text = source_path.read_text(encoding="utf-8").replace("\t", "    ")
    lines = code_text.splitlines()
    if line_numbers:
        lines = [f"{line_no:04}: {line}" if line else f"{line_no:04}:" for line_no, line in enumerate(lines, start=1)]

    code_paragraph = document.add_paragraph()
    code_paragraph.style = body_style_name
    code_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code_paragraph.paragraph_format.first_line_indent = None
    code_paragraph.paragraph_format.line_spacing = 1
    run = code_paragraph.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    rpr = run._element.get_or_add_rPr()
    if rpr.rFonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rpr.rFonts.set(qn("w:ascii"), "Courier New")
    rpr.rFonts.set(qn("w:hAnsi"), "Courier New")
    rpr.rFonts.set(qn("w:eastAsia"), "Courier New")
    rpr.rFonts.set(qn("w:cs"), "Courier New")

    ensure_parent(target_path)
    document.save(target_path)
    return target_path


def summarize_docx(docx_path: Path) -> dict[str, Any]:
    document = Document(str(docx_path))
    non_empty_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    styles_used = sorted({paragraph.style.name for paragraph in non_empty_paragraphs})
    return {
        "document_path": str(docx_path),
        "sections": len(document.sections),
        "paragraphs": len(document.paragraphs),
        "non_empty_paragraphs": len(non_empty_paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "styles_used": styles_used,
    }


def build_sample_vkr_document(
    template_path: Path,
    output_path: Path,
    sample_case: dict[str, Any],
    spec_path: Path | None,
    *,
    macro_names: list[str] | None = None,
    insert_placeholder_captions: bool = False,
    strip_heading_numbering: bool = False,
) -> Path:
    spec = load_spec(spec_path)
    create_docx_from_template(template_path, output_path, macro_names=macro_names)
    apply_gost_profile(output_path, None, spec_path)

    document = Document(str(output_path))
    clear_document(document)

    for section in document.sections:
        set_default_section(section, spec)

    style_map = spec["template_style_map"]
    body_style = get_style_name(document, style_map["body"], "Normal")
    chapter_style = get_style_name(document, style_map["chapter_title"], body_style)
    section_style = get_style_name(document, style_map["section_title"], body_style)
    subsection_style = get_style_name(document, style_map["subsection_title"], body_style)
    toc_title_style = get_style_name(document, style_map["toc_heading"], chapter_style)

    add_blank(document, body_style, 2)
    add_paragraph(
        document,
        body_style,
        "[НАЗВАНИЕ ВУЗА]",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
    )
    add_paragraph(
        document,
        body_style,
        "[ИНСТИТУТ / ФАКУЛЬТЕТ]",
        alignment="center",
        first_line_indent_cm=None,
    )
    add_paragraph(
        document,
        body_style,
        "[КАФЕДРА / ОТДЕЛЕНИЕ]",
        alignment="center",
        first_line_indent_cm=None,
    )
    add_blank(document, body_style, 4)
    add_paragraph(
        document,
        body_style,
        "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=16,
    )
    add_paragraph(
        document,
        body_style,
        "Тестовый образец пояснительной записки",
        alignment="center",
        first_line_indent_cm=None,
        font_size_pt=14,
    )
    add_blank(document, body_style, 2)
    add_paragraph(
        document,
        body_style,
        f"Тема: {sample_case['topic']}",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
    )
    add_blank(document, body_style, 6)
    add_paragraph(document, body_style, "Автор: [ФИО АВТОРА]", alignment="right", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Руководитель: [ФИО РУКОВОДИТЕЛЯ]", alignment="right", first_line_indent_cm=None)
    add_blank(document, body_style, 4)
    add_paragraph(document, body_style, "[ГОРОД] [ГОД]", alignment="center", first_line_indent_cm=None)

    toc_section = document.add_section(WD_SECTION.NEW_PAGE)
    set_default_section(toc_section, spec)
    add_paragraph(
        document,
        toc_title_style,
        spec["toc_rules"]["page_title"],
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
    )
    toc_paragraph = add_paragraph(document, body_style, alignment="left", first_line_indent_cm=None)
    add_field(toc_paragraph, build_toc_instruction(spec), "Оглавление обновится после обновления полей.")

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    set_default_section(body_section, spec)
    add_heading(document, chapter_style, "РЕФЕРАТ", outline_level=0, alignment="center")
    add_paragraph(
        document,
        body_style,
        (
            f"Работа посвящена теме «{sample_case['topic']}». "
            f"Цель исследования состоит в том, чтобы обеспечить {sample_case['goal']}. "
            f"Объект исследования - {sample_case['object']}. "
            f"Предмет исследования - {sample_case['subject']}."
        ),
    )
    add_paragraph(
        document,
        body_style,
        (
            f"В работе рассмотрены особенности предметной области: {sample_case['domain']}. "
            "Предложено программное решение со структурой, достаточной для проверки оформления, оглавления, "
            "иерархии заголовков и списков внутри разделов."
        ),
    )
    add_paragraph(
        document,
        body_style,
        "Ключевые слова: информационная система, проектирование, автоматизация, данные, интерфейс, тестовый документ.",
        first_line_indent_cm=None,
    )

    add_heading(document, chapter_style, "ВВЕДЕНИЕ", outline_level=0, alignment="center", page_break=True)
    add_paragraph(
        document,
        body_style,
        (
            f"Актуальность темы обусловлена тем, что {sample_case['domain']} требует формализованного и "
            "воспроизводимого подхода к сбору, обработке и интерпретации данных. "
            "При отсутствии единого программного решения возрастает трудоемкость сопровождения процессов и снижается "
            "прозрачность управленческих решений."
        ),
    )
    add_paragraph(
        document,
        body_style,
        f"Целью работы является {sample_case['goal']}. Для достижения цели в документе последовательно рассматриваются аналитический, проектный и практический этапы.",
    )
    add_paragraph(document, body_style, "Основные задачи исследования:", first_line_indent_cm=None, bold=True)
    add_list_item(document, body_style, f"проанализировать {sample_case['domain']};")
    add_list_item(document, body_style, "сформировать функциональные и информационные требования к системе;")
    add_list_item(document, body_style, "спроектировать архитектуру, состав модулей и структуру данных;")
    add_list_item(document, body_style, "описать сценарии использования, критерии приемки и ожидаемый эффект внедрения;")

    add_heading(
        document,
        chapter_style,
        "1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ И ПОСТАНОВКА ЗАДАЧИ",
        outline_level=0,
        alignment="center",
        page_break=True,
    )
    add_heading(
        document,
        section_style,
        "1.1 Характеристика предметной области",
        outline_level=1,
        alignment="left",
    )
    add_paragraph(
        document,
        body_style,
        (
            f"Предметная область охватывает {sample_case['domain']}. "
            "В существующей практике существенная часть операций выполняется вручную, что затрудняет контроль сроков, "
            "подготовку отчетности и единообразное представление данных для разных участников процесса."
        ),
    )
    add_paragraph(
        document,
        body_style,
        (
            "Для корректной постановки задачи важно описать участников процесса, информационные объекты, события и "
            "ограничения, влияющие на архитектуру будущего решения. "
            "На этой основе формируется набор требований и критериев качества."
        ),
    )

    add_heading(
        document,
        subsection_style,
        "1.1.1 Проблемы текущего процесса",
        outline_level=2,
        alignment="left",
    )
    add_paragraph(
        document,
        body_style,
        "Анализ текущей практики показывает наличие типовых затруднений, которые препятствуют стабильному выполнению процесса и затрудняют управленческий контроль.",
    )
    for issue in sample_case["pain_points"]:
        add_list_item(document, body_style, issue, level=1)

    add_heading(
        document,
        section_style,
        "1.2 Формирование требований к системе",
        outline_level=1,
        alignment="left",
    )
    add_paragraph(
        document,
        body_style,
        (
            "Требования определяются с учетом ролей пользователей, структуры данных и предполагаемого порядка эксплуатации. "
            "Набор требований должен обеспечивать не только функциональную полноту, но и возможность дальнейшего развития решения."
        ),
    )
    for requirement in sample_case["requirements"]:
        add_list_item(document, body_style, requirement, level=1)

    add_heading(
        document,
        chapter_style,
        "2 ПРОЕКТИРОВАНИЕ АРХИТЕКТУРЫ И ИНФОРМАЦИОННОЙ МОДЕЛИ",
        outline_level=0,
        alignment="center",
        page_break=True,
    )
    add_heading(
        document,
        section_style,
        "2.1 Архитектурная модель решения",
        outline_level=1,
        alignment="left",
    )
    add_paragraph(
        document,
        body_style,
        (
            "Архитектура проектируемой системы принимается модульной. Такой подход позволяет разделить ответственность "
            "между подсистемами, сократить связность компонентов и упростить масштабирование отдельных функций."
        ),
    )
    add_paragraph(document, body_style, "В состав проектируемого решения входят следующие ключевые модули:", first_line_indent_cm=None, bold=True)
    for module_name in sample_case["modules"]:
        add_list_item(document, body_style, module_name, level=1)

    add_heading(
        document,
        subsection_style,
        "2.1.1 Структура данных и сущности",
        outline_level=2,
        alignment="left",
    )
    add_paragraph(
        document,
        body_style,
        (
            "Информационная модель должна поддерживать устойчивые связи между сущностями и позволять отслеживать историю изменений. "
            "Ключевые сущности модели выбираются на основе объектов учета и управленческих сценариев."
        ),
    )
    add_paragraph(document, body_style, "Базовый набор сущностей тестового решения включает:", first_line_indent_cm=None, bold=True)
    for entity in sample_case["entities"]:
        add_list_item(document, body_style, entity, level=1)

    add_heading(
        document,
        section_style,
        "2.2 Проектирование пользовательских сценариев",
        outline_level=1,
        alignment="left",
    )
    add_paragraph(
        document,
        body_style,
        (
            "Пользовательские сценарии описывают последовательность действий, которая приводит к достижению прикладного результата. "
            "Для тестового документа важно отразить как основные, так и контрольные сценарии, участвующие в формировании структуры оглавления."
        ),
    )
    add_paragraph(document, body_style, "В проектируемом решении выделяются следующие сценарии:", first_line_indent_cm=None, bold=True)
    add_list_item(document, body_style, "регистрация исходных данных и их валидация;", level=1)
    add_list_item(document, body_style, "обработка данных и фиксация статуса выполнения операций;", level=1)
    add_list_item(document, body_style, "формирование аналитической информации и итоговых отчетов;", level=1)

    add_heading(
        document,
        chapter_style,
        "3 РЕАЛИЗАЦИЯ, ТЕСТИРОВАНИЕ И ОЦЕНКА РЕЗУЛЬТАТА",
        outline_level=0,
        alignment="center",
        page_break=True,
    )
    add_heading(
        document,
        section_style,
        "3.1 Сценарии реализации и проверки",
        outline_level=1,
        alignment="left",
    )
    add_paragraph(
        document,
        body_style,
        (
            "Реализация тестового решения ориентирована на последовательную сборку модулей, настройку маршрутов обработки данных и "
            "проверку корректности пользовательских операций. Проверка выполняется по набору типовых сценариев и ожидаемых результатов."
        ),
    )
    add_paragraph(document, body_style, "Критерии качества и приемки включают:", first_line_indent_cm=None, bold=True)
    for metric in sample_case["quality_metrics"]:
        add_list_item(document, body_style, metric, level=1)

    add_heading(
        document,
        subsection_style,
        "3.1.1 Тестовые сценарии и критерии приемки",
        outline_level=2,
        alignment="left",
    )
    add_paragraph(
        document,
        body_style,
        (
            "Для подтверждения работоспособности системы задаются сценарии, которые покрывают ввод данных, штатную обработку, "
            "обработку исключений и формирование отчетности. Каждый сценарий должен иметь ожидаемый результат и понятный критерий завершения."
        ),
    )
    add_list_item(document, body_style, "сценарий первичной регистрации данных и проверки обязательных полей;", level=1)
    add_list_item(document, body_style, "сценарий изменения статуса объекта и фиксации истории действий;", level=1)
    add_list_item(document, body_style, "сценарий получения итогового отчета и сравнения контрольных показателей;", level=1)

    add_heading(
        document,
        section_style,
        "3.2 Практический эффект внедрения",
        outline_level=1,
        alignment="left",
    )
    add_paragraph(
        document,
        body_style,
        (
            "Оценка результата ориентирована на сопоставление текущего и целевого состояния процесса. "
            "В тестовом документе акцент делается на тех эффектах, которые позволяют проверить оформление абзацев, списков и выводов в конце раздела."
        ),
    )
    for effect in sample_case["effects"]:
        add_list_item(document, body_style, effect, level=1)

    add_heading(document, chapter_style, "ЗАКЛЮЧЕНИЕ", outline_level=0, alignment="center", page_break=True)
    add_paragraph(
        document,
        body_style,
        (
            f"В ходе подготовки тестовой работы по теме «{sample_case['topic']}» сформирована структура документа, "
            "включающая введение, главы, разделы второго уровня, подпункты третьего уровня, списки и итоговые разделы. "
            "Такой набор элементов позволяет проверить корректность применения ГОСТ-стилей и обновления оглавления."
        ),
    )
    add_paragraph(
        document,
        body_style,
        (
            f"Предложенное решение ориентировано на {sample_case['goal']}. "
            "Содержательная часть документа остается демонстрационной, однако ее объема достаточно для тестирования верстки и логики формирования структуры."
        ),
    )

    add_heading(
        document,
        chapter_style,
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        outline_level=0,
        alignment="center",
        page_break=True,
    )
    for index, source in enumerate(sample_case["sources"], start=1):
        add_paragraph(document, body_style, f"{index}. {source}.", alignment="left", first_line_indent_cm=None)

    add_heading(document, chapter_style, "ПРИЛОЖЕНИЕ А", outline_level=0, alignment="center", page_break=True)
    add_paragraph(document, body_style, "Состав тестовых материалов", alignment="center", first_line_indent_cm=None, bold=True)
    add_paragraph(
        document,
        body_style,
        "В приложение вынесены дополнительные материалы, позволяющие проверить оформление заголовка приложения и перечней вспомогательных элементов.",
    )
    for appendix_item in sample_case["appendix_items"]:
        add_list_item(document, body_style, appendix_item, level=1)

    configure_page_numbering_sections(document, spec)
    ensure_parent(output_path)
    document.save(output_path)
    refresh_fields(
        output_path,
        macro_names=macro_names,
        spec_path=spec_path,
        builtin_layout=True,
        insert_placeholder_captions=insert_placeholder_captions,
        strip_heading_numbering=strip_heading_numbering,
    )
    return output_path


def generate_sample_vkr_set(
    template_path: Path,
    output_dir: Path,
    spec_path: Path | None,
    *,
    include_pdf: bool = False,
    macro_names: list[str] | None = None,
    insert_placeholder_captions: bool = False,
    strip_heading_numbering: bool = False,
) -> dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    generated_items: list[dict[str, str]] = []

    for index, sample_case in enumerate(SAMPLE_VKR_CASES, start=1):
        base_name = f"{index:02d}_{sanitize_filename_component(sample_case['topic'])}"
        docx_path = output_dir / f"{base_name}.docx"
        build_sample_vkr_document(
            template_path,
            docx_path,
            sample_case,
            spec_path,
            macro_names=macro_names,
            insert_placeholder_captions=insert_placeholder_captions,
            strip_heading_numbering=strip_heading_numbering,
        )

        item = {
            "topic": sample_case["topic"],
            "docx": str(docx_path),
        }
        if include_pdf:
            pdf_path = output_dir / f"{base_name}.pdf"
            export_pdf(
                docx_path,
                pdf_path,
                macro_names=macro_names,
                spec_path=spec_path,
                builtin_layout=True,
                insert_placeholder_captions=insert_placeholder_captions,
                strip_heading_numbering=strip_heading_numbering,
            )
            item["pdf"] = str(pdf_path)
        generated_items.append(item)

    return {
        "template": str(template_path),
        "output_dir": str(output_dir),
        "profile": "masters_vkr_vo",
        "documents": generated_items,
    }


def build_agent_brief_document(
    template_path: Path,
    output_path: Path,
    spec_path: Path | None,
    *,
    macro_names: list[str] | None = None,
) -> Path:
    spec = load_spec(spec_path)
    create_docx_from_template(template_path, output_path, macro_names=macro_names)
    document = Document(str(output_path))
    clear_document(document)

    style_map = spec["template_style_map"]
    body_style = get_style_name(document, style_map["body"], "Normal")
    chapter_style = body_style
    section_style = body_style

    apply_gost_profile(output_path, None, spec_path)
    document = Document(str(output_path))
    clear_document(document)

    for section in document.sections:
        set_default_section(section, spec)

    add_blank(document, body_style, 2)
    add_paragraph(
        document,
        chapter_style,
        "ПРОФИЛЬ ДЛЯ ИИ-АГЕНТА ПО ПОДГОТОВКЕ ВКР И ДИПЛОМНЫХ ПРОЕКТОВ",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=16,
        outline_level=0,
    )
    add_paragraph(
        document,
        body_style,
        "Редактируемый Word-образец на базе пользовательского шаблона и ГОСТ-профиля.",
        alignment="center",
        first_line_indent_cm=None,
        font_size_pt=14,
    )
    add_blank(document, body_style, 3)
    add_paragraph(document, body_style, "Тема работы: [ТЕМА РАБОТЫ]", alignment="left", first_line_indent_cm=None, bold=True)
    add_paragraph(document, body_style, "Тип работы: [ВКР / ДИПЛОМНЫЙ ПРОЕКТ / ПОЯСНИТЕЛЬНАЯ ЗАПИСКА]", alignment="left", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Образовательная организация: [НАЗВАНИЕ ВУЗА / КОЛЛЕДЖА]", alignment="left", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Институт / факультет: [ИНСТИТУТ / ФАКУЛЬТЕТ]", alignment="left", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Кафедра / отделение: [КАФЕДРА / ОТДЕЛЕНИЕ]", alignment="left", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Автор: [ФИО АВТОРА]", alignment="left", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Руководитель: [ФИО РУКОВОДИТЕЛЯ]", alignment="left", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Группа: [ГРУППА]", alignment="left", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Город: [ГОРОД]", alignment="left", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Год: [ГОД]", alignment="left", first_line_indent_cm=None)

    document.add_page_break()
    add_paragraph(
        document,
        chapter_style,
        spec["toc_rules"]["page_title"],
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=0,
    )
    toc_paragraph = add_paragraph(document, body_style, alignment="left", first_line_indent_cm=None)
    add_field(toc_paragraph, r'TOC \o "1-2" \h \z \u', "Оглавление обновится после обновления полей.")

    document.add_page_break()
    add_paragraph(
        document,
        chapter_style,
        "1. НАЗНАЧЕНИЕ ДОКУМЕНТА",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=0,
    )
    add_paragraph(document, body_style, spec["description"], font_size_pt=14)
    add_paragraph(document, body_style, spec["source_assets"]["mandatory_usage_rule"], font_size_pt=14)
    add_paragraph(document, body_style, spec["automation_tooling"]["agent_usage_rule"], font_size_pt=14)

    routing = spec["work_type_routing"]
    add_paragraph(
        document,
        section_style,
        "1.1 Как выбрать профиль методических указаний",
        alignment="left",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=1,
    )
    for item in routing["selection_priority"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)
    for rule in routing["ambiguity_rules"]:
        add_paragraph(document, body_style, f"- {rule['trigger']} {rule['action']}", first_line_indent_cm=None)
    for profile in routing["profiles"]:
        blueprint = profile["document_blueprint"]
        add_paragraph(document, body_style, profile["title"], alignment="left", first_line_indent_cm=None, bold=True)
        add_paragraph(
            document,
            body_style,
            (
                f"Основная методичка: {profile['guideline_document']['display_name']} "
                f"({profile['guideline_document']['file_name_hint']})."
            ),
            first_line_indent_cm=None,
        )
        add_paragraph(
            document,
            body_style,
            f"Когда использовать: {'; '.join(profile['when_to_use'])}",
            first_line_indent_cm=None,
        )
        add_paragraph(
            document,
            body_style,
            f"Фронт-маттер: {'; '.join(blueprint['front_matter'])}",
            first_line_indent_cm=None,
        )
        add_paragraph(
            document,
            body_style,
            f"Основные разделы: {'; '.join(blueprint['body'])}",
            first_line_indent_cm=None,
        )
        if "final_sections" in blueprint:
            add_paragraph(
                document,
                body_style,
                f"Завершающие элементы: {'; '.join(blueprint['final_sections'])}",
                first_line_indent_cm=None,
            )
        if "source_minimum" in blueprint:
            add_paragraph(
                document,
                body_style,
                f"Требование к источникам: {blueprint['source_minimum']}.",
                first_line_indent_cm=None,
            )

    add_paragraph(
        document,
        section_style,
        "1.2 Что пользователь должен заполнить перед отправкой ИИ",
        alignment="left",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=1,
    )
    for field_name, description in spec["required_input_fields"].items():
        add_paragraph(document, body_style, f"{field_name}: {description}", first_line_indent_cm=None)

    add_paragraph(
        document,
        section_style,
        "1.3 Политика плейсхолдеров и запроса реквизитов",
        alignment="left",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=1,
    )
    placeholder_policy = spec["input_collection_policy"]
    add_paragraph(document, body_style, placeholder_policy["no_hardcoded_identity_rule"], first_line_indent_cm=None)
    add_paragraph(document, body_style, placeholder_policy["generic_placeholder_format"], first_line_indent_cm=None)
    add_paragraph(document, body_style, placeholder_policy["final_submission_rule"], first_line_indent_cm=None)
    add_paragraph(document, body_style, "Типовые плейсхолдеры для титульных и сопроводительных листов:", first_line_indent_cm=None, bold=True)
    for item in placeholder_policy["title_page_placeholder_examples"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Если отсутствуют обязательные реквизиты, которые нужно запросить перед финальной генерацией:", first_line_indent_cm=None, bold=True)
    for item in placeholder_policy["request_when_missing"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)

    add_paragraph(
        document,
        chapter_style,
        "2. ОБЯЗАТЕЛЬНЫЕ ТРЕБОВАНИЯ ОФОРМЛЕНИЯ",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=0,
    )
    formatting = spec["global_formatting_rules"]
    add_paragraph(
        document,
        body_style,
        (
            f"Основной текст оформляется шрифтом {formatting['main_text']['font_name']} "
            f"{formatting['main_text']['font_size_pt']} pt, межстрочный интервал "
            f"{formatting['main_text']['line_spacing']}, абзацный отступ "
            f"{formatting['main_text']['first_line_indent_cm']} см, выравнивание по ширине."
        ),
    )
    add_paragraph(
        document,
        body_style,
        (
            f"Поля страницы: левое {formatting['margins_mm']['left']} мм, правое "
            f"{formatting['margins_mm']['right']} мм, верхнее {formatting['margins_mm']['top']} мм, "
            f"нижнее {formatting['margins_mm']['bottom']} мм."
        ),
    )
    numbering = formatting["page_numbering"]
    add_paragraph(
        document,
        body_style,
        (
            f"Нумерация страниц арабская, положение {numbering['position']}, печатать номер с раздела "
            f"«{numbering['print_from_section']}». Титульный лист и содержание входят в общий счет, "
            f"но номер на них не печатается."
        ),
    )

    add_paragraph(
        document,
        section_style,
        "2.1 Оглавление",
        alignment="left",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=1,
    )
    toc_rules = spec["toc_rules"]
    add_paragraph(
        document,
        body_style,
        (
            f"Заголовок страницы оглавления и сами записи должны быть оформлены шрифтом "
            f"{toc_rules['entry_format']['font_name']} {toc_rules['entry_format']['font_size_pt']} pt "
            f"с межстрочным интервалом {toc_rules['entry_format']['line_spacing']}."
        ),
    )
    for rule in toc_rules["must_include"]:
        add_paragraph(document, body_style, f"- {rule}", first_line_indent_cm=None)
    for rule in toc_rules["must_not"]:
        add_paragraph(document, body_style, f"- {rule}", first_line_indent_cm=None)

    add_paragraph(
        document,
        chapter_style,
        "3. РЕКОМЕНДУЕМАЯ СТРУКТУРА РАБОТЫ",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=0,
    )
    for item in spec["content_structure_rules"]["minimum_required_sections"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)
    add_paragraph(
        document,
        section_style,
        "3.1 Общие структурные элементы",
        alignment="left",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=1,
    )
    for item in spec["content_structure_rules"]["introduction_must_include"]:
        add_paragraph(document, body_style, f"- Во введении предусмотреть: {item}", first_line_indent_cm=None)

    diploma_profile = next(profile for profile in routing["profiles"] if profile["id"] == "diploma_project_spo")
    diploma_blueprint = diploma_profile["document_blueprint"]
    add_paragraph(
        document,
        section_style,
        "3.2 Профиль дипломного проекта СПО",
        alignment="left",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=1,
    )
    add_paragraph(document, body_style, f"Фронт-маттер: {'; '.join(diploma_blueprint['front_matter'])}", first_line_indent_cm=None)
    for item in spec["content_structure_rules"]["recommended_structure_for_it_project"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)
    add_paragraph(
        document,
        body_style,
        f"Завершающие разделы: {'; '.join(diploma_blueprint['final_sections'])}.",
        first_line_indent_cm=None,
    )
    add_paragraph(
        document,
        body_style,
        (
            f"Рекомендуемый объем основной части: {diploma_blueprint['main_volume_pages']}; "
            f"объем введения: {diploma_blueprint['introduction_volume_pages']}."
        ),
        first_line_indent_cm=None,
    )
    for item in diploma_profile["special_notes"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)

    masters_profile = next(profile for profile in routing["profiles"] if profile["id"] == "masters_vkr_vo")
    masters_blueprint = masters_profile["document_blueprint"]
    add_paragraph(
        document,
        section_style,
        "3.3 Профиль ВКР магистратуры",
        alignment="left",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=1,
    )
    add_paragraph(document, body_style, f"Фронт-маттер: {'; '.join(masters_blueprint['front_matter'])}", first_line_indent_cm=None)
    for item in masters_blueprint["body"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)
    add_paragraph(document, body_style, f"Требование к источникам: {masters_blueprint['source_minimum']}.", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Сопроводительные документы:", first_line_indent_cm=None, bold=True)
    for item in masters_profile["supporting_documents"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Профессиональные акценты:", first_line_indent_cm=None, bold=True)
    for item in masters_profile["activity_profiles"]:
        add_paragraph(document, body_style, f"- {item['name']}: {item['focus']}", first_line_indent_cm=None)
    add_paragraph(document, body_style, "Рекомендуемые нотации моделирования:", first_line_indent_cm=None, bold=True)
    for item in masters_profile["modeling_recommendations"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)

    add_paragraph(
        document,
        chapter_style,
        "4. ТАБЛИЦЫ, РИСУНКИ, ПРИЛОЖЕНИЯ И ЛИСТИНГИ",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=0,
    )
    add_paragraph(document, body_style, spec["table_rules"]["reference_rule"])
    add_paragraph(
        document,
        body_style,
        (
            f"Подпись таблицы оформлять как «{spec['table_rules']['caption_layout']['number_line']}», "
            f"слово «Таблица» размещать {spec['table_rules']['caption_layout']['number_alignment']}."
        ),
    )
    add_paragraph(document, body_style, f"Подпись рисунка: {spec['figure_rules']['caption_format']}.")
    add_paragraph(document, body_style, spec["appendix_rules"]["start_rule"])
    add_paragraph(document, body_style, spec["appendix_rules"]["label_rule"])
    add_paragraph(document, body_style, spec["code_listing_rules"]["placement"])

    add_paragraph(
        document,
        chapter_style,
        "5. КОМАНДЫ АВТОМАТИЗАЦИИ WORD",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=0,
    )
    add_paragraph(document, body_style, "Ниже приведены готовые команды для другого ИИ-агента.", first_line_indent_cm=None)
    for command in spec["automation_tooling"]["recommended_commands"]:
        add_paragraph(document, body_style, command, alignment="left", first_line_indent_cm=None)

    add_paragraph(
        document,
        chapter_style,
        "6. ЧЕК-ЛИСТ ПЕРЕД ФИНАЛЬНОЙ СДАЧЕЙ",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=0,
    )
    for item in spec["validation_checklist"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)

    add_paragraph(
        document,
        chapter_style,
        "7. ЗАПРЕТЫ ДЛЯ ИИ",
        alignment="center",
        first_line_indent_cm=None,
        bold=True,
        font_size_pt=14,
        outline_level=0,
    )
    for item in spec["strict_prohibitions"]:
        add_paragraph(document, body_style, f"- {item}", first_line_indent_cm=None)

    ensure_parent(output_path)
    document.save(output_path)
    refresh_fields(output_path, macro_names=macro_names)
    return output_path


def command_inspect_template(args: argparse.Namespace) -> None:
    save_json(inspect_word_package(args.template), args.output)


def command_create_from_template(args: argparse.Namespace) -> None:
    create_docx_from_template(args.template, args.output, macro_names=args.macro)
    print(str(args.output))


def command_refresh_fields(args: argparse.Namespace) -> None:
    refresh_fields(
        args.input,
        macro_names=args.macro,
        spec_path=args.spec,
        builtin_layout=args.builtin_layout,
        insert_placeholder_captions=args.insert_placeholder_captions,
        strip_heading_numbering=args.strip_heading_numbering,
    )
    print(str(args.input))


def command_export_pdf(args: argparse.Namespace) -> None:
    export_pdf(
        args.input,
        args.output,
        macro_names=args.macro,
        spec_path=args.spec,
        builtin_layout=args.builtin_layout,
        insert_placeholder_captions=args.insert_placeholder_captions,
        strip_heading_numbering=args.strip_heading_numbering,
    )
    print(str(args.output))


def command_apply_gost(args: argparse.Namespace) -> None:
    target = apply_gost_profile(args.input, args.output, args.spec)
    print(str(target))


def command_append_code_appendix(args: argparse.Namespace) -> None:
    target = append_code_appendix(
        input_path=args.input,
        source_path=args.source,
        title=args.title,
        appendix_label=args.label,
        output_path=args.output,
        spec_path=args.spec,
        line_numbers=not args.no_line_numbers,
    )
    print(str(target))


def command_summarize_doc(args: argparse.Namespace) -> None:
    save_json(summarize_docx(args.input), args.output)


def command_resolve_profile(args: argparse.Namespace) -> None:
    spec = load_spec(args.spec)
    result = resolve_work_profile(
        spec,
        work_type=args.work_type,
        education_level=args.education_level,
        qualification_level=args.qualification_level,
    )
    save_json(result, args.output)


def command_finalize_doc(args: argparse.Namespace) -> None:
    working_path = apply_gost_profile(args.input, None, args.spec)
    if args.pdf is not None:
        export_pdf(
            working_path,
            args.pdf,
            macro_names=args.macro,
            spec_path=args.spec,
            builtin_layout=True,
            insert_placeholder_captions=args.insert_placeholder_captions,
            strip_heading_numbering=args.strip_heading_numbering,
        )
        print(json.dumps({"docx": str(working_path), "pdf": str(args.pdf)}, ensure_ascii=False, indent=2))
        return
    refresh_fields(
        working_path,
        macro_names=args.macro,
        spec_path=args.spec,
        builtin_layout=True,
        insert_placeholder_captions=args.insert_placeholder_captions,
        strip_heading_numbering=args.strip_heading_numbering,
    )
    print(str(working_path))


def command_generate_agent_brief(args: argparse.Namespace) -> None:
    target = build_agent_brief_document(args.template, args.output, args.spec, macro_names=args.macro)
    if args.pdf is not None:
        export_pdf(target, args.pdf, macro_names=args.macro)
        print(json.dumps({"docx": str(target), "pdf": str(args.pdf)}, ensure_ascii=False, indent=2))
        return
    print(str(target))


def command_generate_sample_vkrs(args: argparse.Namespace) -> None:
    result = generate_sample_vkr_set(
        template_path=args.template,
        output_dir=args.output_dir,
        spec_path=args.spec,
        include_pdf=args.with_pdf,
        macro_names=args.macro,
        insert_placeholder_captions=args.insert_placeholder_captions,
        strip_heading_numbering=args.strip_heading_numbering,
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))


def add_macro_arguments(parser: argparse.ArgumentParser) -> None:
    parser.add_argument(
        "--macro",
        action="append",
        default=[],
        help=(
            "Имя VBA-макроса для явного запуска через Word.Application.Run. "
            "Флаг можно повторять. Работает только для macro-enabled шаблонов и документов."
        ),
    )


def add_builtin_layout_arguments(parser: argparse.ArgumentParser) -> None:
    parser.add_argument(
        "--builtin-layout",
        action="store_true",
        help=(
            "Применить встроенные правила постобработки документа: нормализацию текста, "
            "таблиц, рисунков и специальных заголовков без запуска VBA."
        ),
    )
    parser.add_argument(
        "--insert-placeholder-captions",
        action="store_true",
        help=(
            "При встроенной постобработке добавлять заглушки подписей к таблицам и рисункам, "
            "если подписи отсутствуют."
        ),
    )
    parser.add_argument(
        "--strip-heading-numbering",
        action="store_true",
        help=(
            "При встроенной постобработке снимать автоматическую нумерацию Word у заголовков ГОСТ, "
            "сохраняя текст заголовков. "
            "Спецзаголовки вроде «Введение» и «Заключение» нормализуются всегда."
        ),
    )


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="CLI-инструмент для ИИ-агентов, работающих с Word-документами по шаблону и ГОСТ-профилю.",
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    inspect_parser = subparsers.add_parser(
        "inspect-template",
        help="Извлечь стили и метаданные пакета из .docx/.docm/.dotx/.dotm в JSON.",
    )
    inspect_parser.add_argument("--template", type=Path, required=True)
    inspect_parser.add_argument("--output", type=Path)
    inspect_parser.set_defaults(func=command_inspect_template)

    create_parser = subparsers.add_parser("create-from-template", help="Создать .docx из .dotx/.dotm-шаблона.")
    create_parser.add_argument("--template", type=Path, required=True)
    create_parser.add_argument("--output", type=Path, required=True)
    add_macro_arguments(create_parser)
    create_parser.set_defaults(func=command_create_from_template)

    refresh_parser = subparsers.add_parser("refresh-fields", help="Обновить поля, включая оглавление.")
    refresh_parser.add_argument("--input", type=Path, required=True)
    refresh_parser.add_argument("--spec", type=Path, default=DEFAULT_SPEC_PATH)
    add_macro_arguments(refresh_parser)
    add_builtin_layout_arguments(refresh_parser)
    refresh_parser.set_defaults(func=command_refresh_fields)

    export_parser = subparsers.add_parser("export-pdf", help="Обновить поля и выгрузить PDF.")
    export_parser.add_argument("--input", type=Path, required=True)
    export_parser.add_argument("--output", type=Path, required=True)
    export_parser.add_argument("--spec", type=Path, default=DEFAULT_SPEC_PATH)
    add_macro_arguments(export_parser)
    add_builtin_layout_arguments(export_parser)
    export_parser.set_defaults(func=command_export_pdf)

    gost_parser = subparsers.add_parser("apply-gost", help="Применить ГОСТ-профиль к стилям и секциям документа.")
    gost_parser.add_argument("--input", type=Path, required=True)
    gost_parser.add_argument("--output", type=Path)
    gost_parser.add_argument("--spec", type=Path, default=DEFAULT_SPEC_PATH)
    gost_parser.set_defaults(func=command_apply_gost)

    appendix_parser = subparsers.add_parser(
        "append-code-appendix",
        help="Добавить приложение с листингом кода в конец документа.",
    )
    appendix_parser.add_argument("--input", type=Path, required=True)
    appendix_parser.add_argument("--source", type=Path, required=True)
    appendix_parser.add_argument("--title", required=True)
    appendix_parser.add_argument("--label", required=True)
    appendix_parser.add_argument("--output", type=Path)
    appendix_parser.add_argument("--spec", type=Path, default=DEFAULT_SPEC_PATH)
    appendix_parser.add_argument("--no-line-numbers", action="store_true")
    appendix_parser.set_defaults(func=command_append_code_appendix)

    summarize_parser = subparsers.add_parser("summarize-doc", help="Выдать краткую JSON-сводку по .docx.")
    summarize_parser.add_argument("--input", type=Path, required=True)
    summarize_parser.add_argument("--output", type=Path)
    summarize_parser.set_defaults(func=command_summarize_doc)

    resolve_parser = subparsers.add_parser(
        "resolve-profile",
        help="Определить профиль методических указаний по типу работы и уровню образования.",
    )
    resolve_parser.add_argument("--work-type", type=str, required=True)
    resolve_parser.add_argument("--education-level", type=str)
    resolve_parser.add_argument("--qualification-level", type=str)
    resolve_parser.add_argument("--spec", type=Path, default=DEFAULT_SPEC_PATH)
    resolve_parser.add_argument("--output", type=Path)
    resolve_parser.set_defaults(func=command_resolve_profile)

    finalize_parser = subparsers.add_parser(
        "finalize-doc",
        help="Применить ГОСТ-профиль, обновить поля и при необходимости экспортировать PDF.",
    )
    finalize_parser.add_argument("--input", type=Path, required=True)
    finalize_parser.add_argument("--spec", type=Path, default=DEFAULT_SPEC_PATH)
    finalize_parser.add_argument("--pdf", type=Path)
    add_macro_arguments(finalize_parser)
    finalize_parser.add_argument(
        "--insert-placeholder-captions",
        action="store_true",
        help="Во время финализации добавлять заглушки подписей к таблицам и рисункам при их отсутствии.",
    )
    finalize_parser.add_argument(
        "--strip-heading-numbering",
        action="store_true",
        help="Во время финализации снимать автоматическую нумерацию Word у заголовков ГОСТ.",
    )
    finalize_parser.set_defaults(func=command_finalize_doc)

    brief_parser = subparsers.add_parser(
        "generate-agent-brief",
        help="Создать Word-образец с инструкциями для ИИ на базе шаблона и ГОСТ-профиля.",
    )
    brief_parser.add_argument("--template", type=Path, required=True)
    brief_parser.add_argument("--output", type=Path, required=True)
    brief_parser.add_argument("--spec", type=Path, default=DEFAULT_SPEC_PATH)
    brief_parser.add_argument("--pdf", type=Path)
    add_macro_arguments(brief_parser)
    brief_parser.set_defaults(func=command_generate_agent_brief)

    sample_parser = subparsers.add_parser(
        "generate-sample-vkrs",
        help="Сгенерировать набор тестовых ВКР с оглавлением, списками и заголовками до уровня 1.1.1.",
    )
    sample_parser.add_argument("--template", type=Path, required=True)
    sample_parser.add_argument("--output-dir", type=Path, default=Path("test"))
    sample_parser.add_argument("--spec", type=Path, default=DEFAULT_SPEC_PATH)
    sample_parser.add_argument("--with-pdf", action="store_true")
    add_macro_arguments(sample_parser)
    sample_parser.add_argument(
        "--insert-placeholder-captions",
        action="store_true",
        help="При генерации тестовых ВКР добавлять заглушки подписей к таблицам и рисункам, если они появятся в документе.",
    )
    sample_parser.add_argument(
        "--strip-heading-numbering",
        action="store_true",
        help="При генерации тестовых ВКР снимать автоматическую нумерацию Word у заголовков ГОСТ.",
    )
    sample_parser.set_defaults(func=command_generate_sample_vkrs)

    return parser


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
